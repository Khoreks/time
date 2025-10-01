from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
import docx
from docx.document import Document
from docx.oxml.shared import qn
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
import zipfile
from io import BytesIO
import re

app = FastAPI()

def load_file(file_content: BytesIO) -> Document:
    return docx.Document(file_content)

def extract_content(doc: Document) -> str:
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_images(doc: Document, file_content: BytesIO) -> list:
    # Извлекаем изображения из docx (они находятся в папке word/media/)
    with zipfile.ZipFile(file_content) as docx_zip:
        image_files = [name for name in docx_zip.namelist() if name.startswith('word/media/')]
        images_data = []
        for img in image_files:
            with docx_zip.open(img) as img_file:
                images_data.append({
                    "filename": img,
                    "data": img_file.read()
                })
    return images_data

def extract_hyperlinks(doc: Document) -> list:
    # Извлечение гиперссылок из документа
    rels = doc.part.rels
    hyperlinks = []
    for rel in rels:
        rel_target = rels[rel].target_ref
        if "hyperlink" in rel_target.lower():
            hyperlinks.append(rel_target)
    # Более точное извлечение из параграфов
    links = []
    for paragraph in doc.paragraphs:
        links.extend(get_links_from_paragraph(paragraph))
    return links

def get_links_from_paragraph(paragraph: Paragraph):
    links = []
    part = paragraph._element
    hyperlinks = part.xpath('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    for link in hyperlinks:
        rId = link.get(qn('r:id'))
        if rId:
            target = paragraph.part.rels[rId].target_ref
            text = "".join([t.text for t in link.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})])
            links.append({"text": text, "url": target})
    return links

@app.post("/process_docx")
async def process_docx(file: UploadFile = File(...)):
    contents = await file.read()
    file_stream = BytesIO(contents)

    # Загружаем документ
    doc = load_file(file_stream)

    # Извлекаем текст
    text = extract_content(doc)

    # Извлекаем изображения
    file_stream.seek(0)  # Сбрасываем указатель
    images = extract_images(doc, file_stream)

    # Извлекаем гиперссылки
    hyperlinks = extract_hyperlinks(doc)

    return JSONResponse({
        "text": text,
        "images_count": len(images),
        "images": [img["filename"] for img in images],
        "hyperlinks": hyperlinks
    })
