def parse_document(filepath: str) -> str:
    """
    Возвращает markdown-строку.
    """
    ext = filepath.lower().split('.')[-1]
    if ext == 'docx':
        return parse_docx(filepath)
    elif ext == 'pdf':
        return parse_pdf(filepath)
    else:
        raise ValueError("Unsupported file type")

from docx import Document
from docx.table import Table
from docx.oxml.shared import qn
import os
from pathlib import Path

def parse_docx(filepath: str, image_dir: str = "./images") -> str:
    doc = Document(filepath)
    md_lines = []
    os.makedirs(image_dir, exist_ok=True)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Заголовки
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1])
            md_lines.append(f"{'#' * level} {text}")
        else:
            md_lines.append(text)

        # Если в параграфе есть изображения — извлекаем их
        for run in para.runs:
            for embedded_obj in run.element.xpath('.//a:blip'):
                # Извлечение изображения (нужно использовать python-docx2txt или docx2python)
                # Пока пропускаем, см. ниже

    # Извлечение таблиц
    for table in doc.tables:
        md_lines.append("\n| " + " | ".join([cell.text for cell in table.rows[0].cells]) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |")
        for row in table.rows[1:]:
            md_lines.append("| " + " | ".join([cell.text for cell in row.cells]) + " |")

    return "\n\n".join(md_lines)

import fitz  # PyMuPDF
import pandas as pd

def parse_pdf(filepath: str, image_dir: str = "./images") -> str:
    doc = fitz.open(filepath)
    md_lines = []
    os.makedirs(image_dir, exist_ok=True)

    for page_num in range(len(doc)):
        page = doc[page_num]

        # Извлечение текста
        text = page.get_text("text")
        md_lines.append(text)

        # Извлечение таблиц (если есть, через pandas или pdfplumber)
        # tables = page.find_tables()  # fitz может находить таблицы
        # for table in tables:
        #     df = pd.DataFrame(table)
        #     md_lines.append(df.to_markdown())

        # Извлечение изображений
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            image_filename = f"{image_dir}/page_{page_num}_img_{img_index}.{image_ext}"
            with open(image_filename, "wb") as img_file:
                img_file.write(image_bytes)
            md_lines.append(f"\n![Image](./{image_filename})\n")

    return "\n\n".join(md_lines)
